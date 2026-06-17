from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import html


ROOT = Path(r"C:\Users\HP 250 G10\Documents\GITHUT\MEC\BT\BTCC\Documentos_Generales")
HTML_PATH = ROOT / "Proyecto_Educativo_Competencia_Integrada_BTCC_2026_2031.html"
DOCX_PATH = ROOT / "Proyecto_Educativo_Competencia_Integrada_BTCC_2026_2031.docx"


def fix_text(text: str) -> str:
    if not text:
        return ""
    cleaned = " ".join(text.replace("\xa0", " ").split())
    if any(mark in cleaned for mark in ("Ã", "Â", "â")):
        try:
            cleaned = cleaned.encode("latin1").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            pass
    return cleaned


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def restart_page_numbering(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("BTCC Title", 24, "0055A4"),
        ("BTCC H1", 17, "0055A4"),
        ("BTCC H2", 13, "000000"),
    ):
        if name not in document.styles:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = document.styles[name]
        style.base_style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        if name == "BTCC Title":
            style.paragraph_format.space_after = Pt(10)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(5)


def append_inline(paragraph, node, bold=False, italic=False) -> None:
    if node.text and fix_text(node.text):
        run = paragraph.add_run(fix_text(node.text))
        run.bold = bold
        run.italic = italic

    for child in node:
        child_tag = child.tag.lower() if isinstance(child.tag, str) else ""
        child_bold = bold or child_tag in {"strong", "b", "th"}
        child_italic = italic or child_tag in {"em", "i"}
        append_inline(paragraph, child, child_bold, child_italic)
        if child.tail and fix_text(child.tail):
            tail = paragraph.add_run(fix_text(child.tail))
            tail.bold = bold
            tail.italic = italic


def add_paragraph_from_node(document: Document, node, style=None, indent=0) -> None:
    paragraph = document.add_paragraph(style=style)
    if indent:
        paragraph.paragraph_format.left_indent = Cm(indent)
    append_inline(paragraph, node)


def add_list(document: Document, node, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in node.xpath("./li"):
        paragraph = document.add_paragraph(style=style)
        append_inline(paragraph, item)


def add_callout(document: Document, node) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E5E5E5")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    append_inline(paragraph, node)


def add_table(document: Document, node) -> None:
    rows = node.xpath("./tr")
    if not rows:
        rows = node.xpath("./tbody/tr")
    if not rows:
        return

    first_row = rows[0]
    col_count = len(first_row.xpath("./th|./td"))
    table = document.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, source_row in enumerate(rows):
        cells = source_row.xpath("./th|./td")
        row = table.add_row()
        for idx, source_cell in enumerate(cells):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            append_inline(p, source_cell, bold=(row_index == 0 or source_cell.tag.lower() == "th"))
            if row_index == 0:
                set_cell_shading(cell, "E5E5E5")


def add_badges(document: Document, node) -> None:
    texts = [fix_text(" ".join(child.itertext())) for child in node.xpath(".//span|.//div")]
    texts = [text for text in texts if text]
    if texts:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(" | ".join(texts))
        run.bold = True


def render_sheet(document: Document, section_node) -> None:
    for child in section_node:
        tag = child.tag.lower() if isinstance(child.tag, str) else ""
        classes = set((child.get("class") or "").split())

        if tag == "h2":
            add_paragraph_from_node(document, child, style="BTCC H1")
        elif tag == "h3":
            add_paragraph_from_node(document, child, style="BTCC H2")
        elif tag == "p":
            add_paragraph_from_node(document, child)
        elif tag == "ul":
            add_list(document, child, ordered=False)
        elif tag == "ol":
            add_list(document, child, ordered=True)
        elif tag == "table":
            add_table(document, child)
        elif tag == "div" and "callout" in classes:
            add_callout(document, child)
        elif tag == "div" and "badge-row" in classes:
            add_badges(document, child)
        elif tag == "div":
            for nested in child:
                nested_tag = nested.tag.lower() if isinstance(nested.tag, str) else ""
                if nested_tag == "p":
                    add_paragraph_from_node(document, nested)
                elif nested_tag == "ul":
                    add_list(document, nested, ordered=False)
                elif nested_tag == "ol":
                    add_list(document, nested, ordered=True)
                elif nested_tag == "table":
                    add_table(document, nested)


def build_cover(document: Document, cover_node) -> None:
    title = fix_text(" ".join(cover_node.xpath("./h1//text()")))

    # Collect all p texts
    lines = []
    for paragraph in cover_node.xpath("./p"):
        text = fix_text(" ".join(paragraph.itertext()))
        if text:
            lines.append(text)

    # Meta info from the inline div
    meta_div = cover_node.xpath(".//div[not(@class)]")
    meta_lines = []
    if meta_div:
        for p in meta_div[0].xpath(".//p"):
            text = fix_text(" ".join(p.itertext()))
            if text:
                meta_lines.append(text)

    # "Año 2026" line
    year_text = ""
    for p in cover_node.xpath(".//p"):
        txt = fix_text(" ".join(p.itertext()))
        if "Año" in txt and "2026" in txt:
            year_text = txt

    document.add_paragraph()

    # Header lines (Ministerio, Colegio, BTCC) - first 3 lines
    for i in range(min(3, len(lines))):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(lines[i])
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(2)

    # Title
    p = document.add_paragraph(style="BTCC Title")
    p.paragraph_format.space_before = Pt(30)
    p.add_run(title)

    # OBRATEC 2026 big text (look for the large styled p after h1)
    obratec_big = ""
    for p_el in cover_node.xpath("./p"):
        txt = fix_text(" ".join(p_el.itertext()))
        if "OBRATEC 2026" in txt and "Año" not in txt:
            obratec_big = txt
            break
    if obratec_big:
        document.add_paragraph()
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(obratec_big)
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor.from_string("0055A4")
        p.paragraph_format.space_after = Pt(6)

    document.add_paragraph()

    # Rest of lines (Alcance, Período, Proyectista)
    for line in lines[4:]:
        if "Año" in line and "2026" in line:
            continue
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)

    # Meta info as a compact table
    if meta_lines:
        document.add_paragraph()
        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for ml in meta_lines:
            if ":" in ml:
                parts = ml.split(" ", 1)
                label = parts[0].strip()
                value = parts[1].strip() if len(parts) > 1 else ""
                row = table.add_row()
                left = row.cells[0]
                right = row.cells[1]
                left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_shading(left, "F4F8FD")
                left.paragraphs[0].add_run(label).bold = True
                right.paragraphs[0].add_run(value)

    # Año 2026
    if year_text:
        document.add_paragraph()
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(year_text)
        run.bold = True
        run.font.size = Pt(18)


def main() -> None:
    root = html.fromstring(HTML_PATH.read_text(encoding="utf-8"))
    document = Document()
    configure_document(document)

    cover_node = root.xpath("//section[contains(@class, 'cover')]")[0]
    build_cover(document, cover_node)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    content_section.page_width = Cm(21)
    content_section.page_height = Cm(29.7)
    content_section.top_margin = Cm(2.0)
    content_section.bottom_margin = Cm(1.7)
    content_section.left_margin = Cm(2.0)
    content_section.right_margin = Cm(2.0)
    restart_page_numbering(content_section, start=1)
    content_section.footer.is_linked_to_previous = False
    footer = content_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Página ")
    footer_run.font.name = "Times New Roman"
    footer_run.font.size = Pt(9)
    add_page_number(footer)

    for sheet in root.xpath("//section[contains(@class, 'sheet')]"):
        render_sheet(document, sheet)

    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
