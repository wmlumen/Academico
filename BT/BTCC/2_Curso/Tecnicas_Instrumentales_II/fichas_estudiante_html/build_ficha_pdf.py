from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import html


FICHAS_DIR = Path(__file__).parent
OUT_DIR = FICHAS_DIR / "pdf_compilados"

TEMAS = [
    ("01", "Método de trabajo", "Marzo"),
    ("02", "Organización del puesto y movimiento de materiales", "Abril"),
    ("03", "Estudio de tiempos y muestreo", "Mayo"),
    ("04", "Rendimiento y costo de mano de obra", "Junio"),
    ("05", "Integración de etapa", "Julio"),
    ("06", "Cómputo de trabajos", "Agosto"),
    ("07", "Cómputo de materiales", "Septiembre"),
    ("08", "Cómputo aplicado al hormigón armado", "Octubre"),
    ("09", "Integración final", "Noviembre"),
]

TEMA_PADRE_NAMES = {
    "01": "tema_01_metodo_de_trabajo",
    "02": "tema_02_organizacion_del_puesto",
    "03": "tema_03_estudio_de_tiempos_y_muestreo",
    "04": "tema_04_rendimiento_y_costo_mano_de_obra",
    "05": "tema_05_integracion_de_etapa",
    "06": "tema_06_computo_de_trabajos",
    "07": "tema_07_computo_de_materiales",
    "08": "tema_08_computo_hormigon_armado",
    "09": "tema_09_integracion_final",
}

SUB_FICHA_NAMES = {
    "01": [f"tema_01{a}" for a in "abcd"],
    "02": [f"tema_02{a}" for a in "abcd"],
    "03": [f"tema_03{a}" for a in "abcd"],
    "04": [f"tema_04{a}" for a in "abcd"],
    "05": [f"tema_05{a}" for a in "abc"],
    "06": [f"tema_06{a}" for a in "abcd"],
    "07": [f"tema_07{a}" for a in "abcd"],
    "08": [f"tema_08{a}" for a in "abcd"],
    "09": [f"tema_09{a}" for a in "abcd"],
}

SUB_CLASS_NAMES = {
    "01": ["Introducción al método", "Secuencia y organización", "Aplicación práctica", "Integración y evaluación"],
    "02": ["Principios de organización", "Movimiento de materiales", "Economía de movimientos", "Aplicación práctica"],
    "03": ["Introducción a tiempos", "Muestreo del trabajo", "Tiempos improductivos", "Planillas de registro"],
    "04": ["Jornal, cuadrilla y rendimiento", "Cálculo costo unitario", "Control de producción", "Ejercicios prácticos"],
    "05": ["Repaso general", "Ejercicios integradores", "Evaluación de etapa"],
    "06": ["Introducción al cómputo", "Cómputo albañilería", "Cómputo revoques", "Planilla de cómputo"],
    "07": ["Introducción consumos", "Materiales albañilería", "Materiales pisos", "Ejercicios prácticos"],
    "08": ["Introducción HA", "Volumen y encofrado", "Cómputo acero", "Ejercicios HA"],
    "09": ["Repaso 2da etapa", "Caso integrador", "Planilla técnica", "Evaluación final"],
}


def fix_text(text):
    if not text:
        return ""
    cleaned = " ".join(text.replace("\xa0", " ").split())
    if any(mark in cleaned for mark in ("Ã", "Â", "â")):
        try:
            cleaned = cleaned.encode("latin1").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            pass
    return cleaned


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def restart_page_numbering(section, start=1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def configure_document(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("BTCC Title", 24, "0055A4"),
        ("BTCC H1", 17, "0055A4"),
        ("BTCC H2", 13, "000000"),
    ):
        if name not in document.styles:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = document.styles[name]
        style.base_style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        if name == "BTCC Title":
            style.paragraph_format.space_after = Pt(10)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(5)


def append_inline(paragraph, node, bold=False, italic=False):
    if node.text and fix_text(node.text):
        run = paragraph.add_run(fix_text(node.text))
        run.bold = bold
        run.italic = italic

    for child in node:
        child_tag = child.tag.lower() if isinstance(child.tag, str) else ""
        child_bold = bold or child_tag in {"strong", "b", "th"}
        child_italic = italic or child_tag in {"em", "i"}
        append_inline(paragraph, child, child_bold, child_italic)
        if child.tail and fix_text(child.tail):
            tail = paragraph.add_run(fix_text(child.tail))
            tail.bold = bold
            tail.italic = italic


def add_para(document, node, style=None, indent=0):
    paragraph = document.add_paragraph(style=style)
    if indent:
        paragraph.paragraph_format.left_indent = Cm(indent)
    append_inline(paragraph, node)


def add_list(document, node, ordered):
    style = "List Number" if ordered else "List Bullet"
    for item in node.xpath("./li"):
        paragraph = document.add_paragraph(style=style)
        append_inline(paragraph, item)


def add_callout(document, node):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E5E5E5")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    append_inline(paragraph, node)


def add_table(document, node):
    rows = node.xpath("./tr")
    if not rows:
        rows = node.xpath("./tbody/tr")
    if not rows:
        return

    first_row = rows[0]
    col_count = len(first_row.xpath("./th|./td"))
    table = document.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, source_row in enumerate(rows):
        cells = source_row.xpath("./th|./td")
        row = table.add_row()
        for idx, source_cell in enumerate(cells):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            append_inline(p, source_cell, bold=(row_index == 0 or source_cell.tag.lower() == "th"))
            if row_index == 0:
                set_cell_shading(cell, "E5E5E5")


def render_sheet(document, section_node):
    for child in section_node:
        tag = child.tag.lower() if isinstance(child.tag, str) else ""
        classes = set((child.get("class") or "").split())

        if tag == "h1":
            add_para(document, child, style="BTCC H1")
        elif tag == "h2":
            add_para(document, child, style="BTCC H1")
        elif tag == "h3":
            add_para(document, child, style="BTCC H2")
        elif tag == "p":
            add_para(document, child)
        elif tag == "ul":
            add_list(document, child, ordered=False)
        elif tag == "ol":
            add_list(document, child, ordered=True)
        elif tag == "table":
            add_table(document, child)
        elif tag == "div" and "callout" in classes:
            add_callout(document, child)
        elif tag == "div" and "badge-row" in classes:
            pass
        elif tag == "div" and "firma" in classes:
            pass
        elif tag == "div":
            for nested in child:
                nested_tag = nested.tag.lower() if isinstance(nested.tag, str) else ""
                if nested_tag == "p":
                    add_para(document, nested)
                elif nested_tag == "ul":
                    add_list(document, nested, ordered=False)
                elif nested_tag == "ol":
                    add_list(document, nested, ordered=True)
                elif nested_tag == "table":
                    add_table(document, nested)


def build_cover(document, tema_num, tema_name, mes):
    document.add_paragraph()

    for line in ["Ministerio de Educación y Ciencias", "Colegio Nacional Defensores del Chaco"]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(13)
        p.paragraph_format.space_after = Pt(4)

    document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Técnicas Instrumentales II")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("0055A4")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Segundo Curso – BTCC")
    run.font.size = Pt(12)

    document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Material del Estudiante")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0055A4")
    p.paragraph_format.space_after = Pt(20)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"TEMA {tema_num} – {tema_name.upper()}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("000000")
    p.paragraph_format.space_after = Pt(4)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Mes: {mes}  |  16 h cátedra")
    run.font.size = Pt(11)

    document.add_paragraph()

    sub_names = SUB_CLASS_NAMES[tema_num]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, sn in enumerate(sub_names):
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(f"Clase {i+1}").bold = True
        row.cells[1].paragraphs[0].add_run(sn)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Año 2026")
    run.bold = True
    run.font.size = Pt(14)


def build_tema_pdf(tema_num, tema_name, mes):
    """Build DOCX for a tema compiling its parent + sub-fichas."""
    document = Document()
    configure_document(document)

    # Cover
    build_cover(document, tema_num, tema_name, mes)

    # Parent ficha
    parent_file = FICHAS_DIR / f"{TEMA_PADRE_NAMES[tema_num]}.html"
    if parent_file.exists():
        content_section = document.add_section(WD_SECTION.NEW_PAGE)
        content_section.page_width = Cm(21)
        content_section.page_height = Cm(29.7)
        content_section.top_margin = Cm(2.0)
        content_section.bottom_margin = Cm(1.7)
        content_section.left_margin = Cm(2.0)
        content_section.right_margin = Cm(2.0)
        content_section.footer.is_linked_to_previous = False
        footer = content_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer.add_run("Página ")
        fr.font.name = "Times New Roman"
        fr.font.size = Pt(9)
        add_page_number(footer)

        root = html.fromstring(parent_file.read_text(encoding="utf-8"))
        for sheet in root.xpath("//section[contains(@class, 'sheet')]"):
            render_sheet(document, sheet)
        for sheet in root.xpath("//section[not(contains(@class, 'sheet'))]"):
            if sheet.tag.lower() == "section":
                render_sheet(document, sheet)

    # Sub-fichas
    sub_files = SUB_FICHA_NAMES.get(tema_num, [])
    for sf in sub_files:
        sf_file = FICHAS_DIR / f"{sf}.html"
        if not sf_file.exists():
            continue
        content_section = document.add_section(WD_SECTION.NEW_PAGE)
        content_section.page_width = Cm(21)
        content_section.page_height = Cm(29.7)
        content_section.top_margin = Cm(2.0)
        content_section.bottom_margin = Cm(1.7)
        content_section.left_margin = Cm(2.0)
        content_section.right_margin = Cm(2.0)
        content_section.footer.is_linked_to_previous = False
        footer = content_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer.add_run("Página ")
        fr.font.name = "Times New Roman"
        fr.font.size = Pt(9)
        add_page_number(footer)

        root = html.fromstring(sf_file.read_text(encoding="utf-8"))
        for sheet in root.xpath("//section[contains(@class, 'sheet')]"):
            render_sheet(document, sheet)
        for sheet in root.xpath("//section[not(contains(@class, 'sheet'))]"):
            if sheet.tag.lower() == "section":
                render_sheet(document, sheet)

    # Save DOCX
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUT_DIR / f"Tema_{tema_num}_Material_del_Estudiante.docx"
    document.save(str(docx_path))
    return docx_path


def docx_to_pdf(docx_path):
    """Convert DOCX to PDF using Word COM automation."""
    import subprocess
    ps_script = f'''
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $doc = $word.Documents.Open("{docx_path}")
    $pdfPath = "{docx_path}".Replace(".docx", ".pdf")
    $doc.SaveAs([ref] $pdfPath, [ref] 17)
    $doc.Close()
    $word.Quit()
    '''
    try:
        subprocess.run(["powershell", "-Command", ps_script], check=True, capture_output=True, timeout=60)
        return str(docx_path).replace(".docx", ".pdf")
    except subprocess.CalledProcessError as e:
        print(f"  Word COM error: {e.stderr.decode() if e.stderr else 'unknown'}")
        return None


def main():
    print("Building compilation PDFs for all 9 temas...")
    for tema_num, tema_name, mes in TEMAS:
        print(f"  Tema {tema_num}: {tema_name}...")
        docx_path = build_tema_pdf(tema_num, tema_name, mes)
        print(f"    DOCX: {docx_path}")
        pdf_path = docx_to_pdf(docx_path)
        if pdf_path:
            print(f"    PDF:  {pdf_path}")
        else:
            print(f"    PDF:  SKIPPED (Word COM unavailable)")
    print("Done.")


if __name__ == "__main__":
    main()
